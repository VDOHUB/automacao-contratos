from docx import Document

for tipo, path in [
    ('empreiteiro', '0\xb0 - Modelos de Contratos/MODELO CONTRATO Empreiteiro - Atualizado.docx'),
    ('estrutura',   '0\xb0 - Modelos de Contratos/MODELO CONTRATO LUIZ CIVIL - Estrutura (1).docx'),
    ('alvenaria',   '0\xb0 - Modelos de Contratos/MODELO CONTRATO LUIZ - Alvenaria 1\xb0PAV (1).docx'),
]:
    doc = Document(path)
    print(f'=== {tipo.upper()} ===')
    print(f'Tabelas: {len(doc.tables)}')
    for ti, tbl in enumerate(doc.tables):
        first = tbl.rows[0].cells[0].text.strip()[:40]
        print(f'  Tabela {ti} ({len(tbl.rows)}x{len(tbl.columns)}): "{first}"')
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if 'Rua' in t or 'Quadra' in t or 'logradouro' in t or 'residente' in t.lower():
            print(f'  Para [{i}] endereco: {repr(t[:120])}')
    print()

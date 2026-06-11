from docx import Document

doc = Document('0\xb0 - Modelos de Contratos/MODELO CONTRATO Empreiteiro - Atualizado.docx')
# Full text of paragraphs 4, 5, 6
for i in range(3, 14):
    p = doc.paragraphs[i]
    print(f'Para [{i}]: {repr(p.text)}')
    print()

print('=== TABELAS ===')
for ti, tbl in enumerate(doc.tables):
    print(f'\nTabela {ti} ({len(tbl.rows)}x{len(tbl.columns)}):')
    for ri, row in enumerate(tbl.rows):
        cells = [repr(c.text.strip()[:80]) for c in row.cells]
        print(f'  Linha {ri}: {cells}')

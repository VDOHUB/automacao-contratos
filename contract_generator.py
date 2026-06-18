"""
Gerador de contratos via docxtpl.
Preprocessa modelos originais (python-docx) substituindo exemplo → {{ variável }}.
"""

import re
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from lxml import etree

BASE_DIR = Path(__file__).parent
MODELOS_DIR = BASE_DIR / "0° - Modelos de Contratos"
TEMPLATES_DIR = BASE_DIR / "contract_templates"
GENERATED_DIR = BASE_DIR / "generated"

TEMPLATES_DIR.mkdir(exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Mapa de contratos disponíveis
# ---------------------------------------------------------------------------

CONTRATOS = {
    "empreiteiro": {
        "label": "Empreiteiro",
        "modelo": "MODELO CONTRATO Empreiteiro - Atualizado.docx",
        "template": "tpl_empreiteiro.docx",
    },
    "alvenaria": {
        "label": "Alvenaria 1° PAV",
        "modelo": "MODELO CONTRATO LUIZ - Alvenaria 1°PAV (1).docx",
        "template": "tpl_alvenaria.docx",
    },
    "arrimo": {
        "label": "Arrimo",
        "modelo": "MODELO CONTRATO LUIZ CIVIL - Arrimo (1).docx",
        "template": "tpl_arrimo.docx",
    },
    "estrutura": {
        "label": "Estrutura Civil",
        "modelo": "MODELO CONTRATO LUIZ CIVIL - Estrutura (1).docx",
        "template": "tpl_estrutura.docx",
    },
    "carpinteiro_1pav": {
        "label": "Carpinteiro 1° PAV",
        "modelo": "MODELO CONTRATO SERGIO - CARPINTEIRO - 1° PAV.docx",
        "template": "tpl_carpinteiro_1pav.docx",
    },
    "carpinteiro_2pav": {
        "label": "Carpinteiro 2° PAV",
        "modelo": "MODELO CONTRATO SERGIO - CARPINTEIRO - 2° PAV.docx",
        "template": "tpl_carpinteiro_2pav.docx",
    },
    "armador": {
        "label": "Armador 1° PAV",
        "modelo": "MODELO CONTRATO WILLIAM CIVIL - ARMADOR - 1° PAV.docx",
        "template": "tpl_armador.docx",
    },
    "piscina": {
        "label": "Piscina",
        "modelo": "MODELO CONTRATO ATUALIZADO - PISCINA.docx",
        "template": "tpl_piscina.docx",
    },
}

# ---------------------------------------------------------------------------
# Substituições por tipo de contrato
# Formato: (texto_exato_no_template, variavel_jinja2)
# ---------------------------------------------------------------------------

# Dados do Contratante — comuns a todos os modelos LUIZ/SERGIO/WILLIAM
_CONTRATANTE_IVONIR = [
    ("Ivonir Vieira dos Santos", "{{ nome_contratante }}"),
    # CPF/RG devem vir antes de partes do texto que os contenham
    ("165.477.301-82", "{{ cpf_contratante }}"),
    ("FV27900 DPF GO", "{{ rg_contratante }}"),
    # Endereço (estrutura, arrimo, armador, carpinteiro_1pav)
    ("Rua GT02, QD 04, LT 11, Residencial Grand Trianon,", "{{ logradouro_contratante }},"),
    # Endereço alternativo (alvenaria)
    ("Rua GT02, QD 04, LT 11, Residencial Grand Trianon, Anápolis-GO,",
     "{{ logradouro_contratante }}, {{ complemento_contratante }},"),
    # Bloco civil — versão mais específica primeiro (carpinteiro_2pav tem 2 gêneros)
    ("brasileiro, brasileira, aposentada",
     "{{ nacionalidade_contratante }}, {{ estado_civil_contratante }}, {{ profissao_contratante }}"),
    ("brasileira, solteira, aposentada",
     "{{ nacionalidade_contratante }}, {{ estado_civil_contratante }}, {{ profissao_contratante }}"),
    # Assinatura em maiúsculo
    ("IVONIR VIEIRA DOS SANTOS", "{{ nome_contratante | upper }}"),
    # Cidade/data assinatura (padrão desses modelos)
    ("Anápolis-GO, 23 de Setembro de 2025.", "{{ cidade_assinatura }}, {{ data_assinatura }}."),
    ("Anápolis-GO, 23 de Setembro de 2025",  "{{ cidade_assinatura }}, {{ data_assinatura }}"),
    ("Anápolis-GO, 07 de Abril de 2026.",    "{{ cidade_assinatura }}, {{ data_assinatura }}."),
    ("Anápolis-GO, 07 de Abril de 2026",     "{{ cidade_assinatura }}, {{ data_assinatura }}"),
    ("Anápolis-GO, 05 de Janeiro de 2026.",  "{{ cidade_assinatura }}, {{ data_assinatura }}."),
    ("Anápolis-GO, 05 de Janeiro de 2026",   "{{ cidade_assinatura }}, {{ data_assinatura }}"),
    ("Anápolis-GO, 09 de Março de 2026.",    "{{ cidade_assinatura }}, {{ data_assinatura }}."),
    ("Anápolis-GO, 09 de Março de 2026",     "{{ cidade_assinatura }}, {{ data_assinatura }}"),
]

_CONTRATANTE_CASSIO = [
    ("Cassio Antônio da Silva", "{{ nome_contratante }}"),
    ("082.911.871-34", "{{ cpf_contratante }}"),
    ("245702", "{{ rg_contratante }}"),
    ("SSP/GO", "SSP/{{ uf_rg_contratante }}"),
    ("brasileiro", "{{ nacionalidade_contratante }}"),
    ("casado", "{{ estado_civil_contratante }}"),
    ("bancário", "{{ profissao_contratante }}"),
    # logradouro — o parágrafo seguinte tem o complemento; serão fundidos depois
    ("Rua Benedito Borges de Almeida,", "{{ logradouro_contratante }},"),
    (
        "Quadra 20, Lote 1, sem número, apartamento 601, Condomínio Forma Opus, Bairro Jundiaí, Cep: 751-10070",
        "{{ complemento_contratante }}",
    ),
    ("CASSIO ANTÔNIO DA SILVA", "{{ nome_contratante | upper }}"),
    ("Anápolis-GO, 22 de Janeiro de 2026", "{{ cidade_assinatura }}, {{ data_assinatura }}"),
]

# Executoras por template
_EXEC_LCJ = [
    ("LCJ Construtora LTDA", "{{ nome_executora }}"),
    ("60.482.939/0001-61", "{{ cnpj_executora }}"),
    ("LCJ CONSTRUTORA LTDA – 60.482.939/0001-61", "{{ nome_executora | upper }} – {{ cnpj_executora }}"),
    ("LCJ CONSTRUTORA LTDA – 60.482.939/0001-61", "{{ nome_executora | upper }} – {{ cnpj_executora }}"),
]

_EXEC_TOP_CARPINTARIA = [
    ("TOP CARPINTARIA LTDA", "{{ nome_executora }}"),
    ("Top Carpintaria LTDA", "{{ nome_executora }}"),
    ("59.049.581/0001-07", "{{ cnpj_executora }}"),
    ("TOP CARPINTARIA LTDA – 59.049.581/0001-07", "{{ nome_executora | upper }} – {{ cnpj_executora }}"),
    ("TOP CARPINTARIA LTDA – 59.049.581/0001-07", "{{ nome_executora | upper }} – {{ cnpj_executora }}"),
]

_EXEC_SOBERANA = [
    ("Soberana Piscinas Ltda", "{{ nome_executora }}"),
    ("SOBERANA PISCINAS LTDA", "{{ nome_executora | upper }}"),
    ("40.244.240/0001-56", "{{ cnpj_executora }}"),
    ("SOBERANA PISCINAS LTDA – 40.244.240/0001-56.", "{{ nome_executora | upper }} – {{ cnpj_executora }}."),
    ("SOBERANA PISCINAS LTDA – 40.244.240/0001-56.", "{{ nome_executora | upper }} – {{ cnpj_executora }}."),
]

_EXEC_MARCOS = [
    ("Marcos Xavier Barbosa", "{{ nome_executora }}"),
    ("034.006.671-70", "{{ cpf_executora }}"),
    ("Rua 1, Quadra 24, lote 12, Residencial Recanto do Sol", "{{ endereco_executora }}"),
    ("MARCOS XAVIER BARBOSA", "{{ nome_executora | upper }}"),
]

# Datas por template
_DATAS = {
    "empreiteiro":      ("06/02/2026", "3 meses", "06/05/2026"),
    "estrutura":        ("22/09/2025", "6 semanas", "31/10/2025"),
    "alvenaria":        ("06/04/2026", "5 semanas", "08/05/2026"),
    "arrimo":           ("22/09/2025", "6 semanas", "31/10/2025"),
    "carpinteiro_1pav": ("05/01/2026", "6 semanas", "13/02/2026"),
    "carpinteiro_2pav": ("09/03/2026", "10 semanas", "18/05/2026"),
    "armador":          ("22/09/2025", "6 semanas", "31/10/2025"),
    "piscina":          ("07/04/2026", "10 semanas", "16/06/2026"),
}

# Multa por template (texto exato antes da vírgula/ponto)
_MULTA = {
    "empreiteiro":      ("R$   3.000,00 (Três mil reais)", "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "estrutura":        ("R$ 8.934,75 (Oito mil novecentos e trinta e quatro reais e setenta e cinco centavos)",
                         "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "alvenaria":        ("R$10.000,00 (dez mil reais)", "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "arrimo":           ("R$ 8.934,75 (Oito mil novecentos e trinta e quatro reais e setenta e cinco centavos)",
                         "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "carpinteiro_1pav": ("R$10.000,00 (dez mil reais)", "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "carpinteiro_2pav": ("R$10.000,00 (dez mil reais)", "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "armador":          ("R$ 8.934,75 (Oito mil novecentos e trinta e quatro reais e setenta e cinco centavos)",
                         "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
    "piscina":          ("R$10.000,00 (dez mil reais)", "{{ valor_multa_fmt }} ({{ valor_multa_extenso }})"),
}


def _get_replacements(tipo: str) -> list[tuple]:
    """Retorna lista completa de (texto_antigo, novo_texto_jinja2) para o tipo."""
    reps = []

    # Datas
    if tipo in _DATAS:
        di, pe, dc = _DATAS[tipo]
        reps += [
            (di, "{{ data_inicio }}"),
            (pe, "{{ prazo_execucao }}"),
            (dc, "{{ data_conclusao }}"),
        ]

    # Multa
    if tipo in _MULTA:
        reps.append(_MULTA[tipo])

    if tipo == "empreiteiro":
        reps += _CONTRATANTE_CASSIO + _EXEC_MARCOS
    elif tipo in ("estrutura", "arrimo", "armador"):
        reps += _CONTRATANTE_IVONIR + _EXEC_LCJ
    elif tipo == "alvenaria":
        reps += _CONTRATANTE_IVONIR + _EXEC_LCJ
    elif tipo in ("carpinteiro_1pav", "carpinteiro_2pav"):
        reps += _CONTRATANTE_IVONIR + _EXEC_TOP_CARPINTARIA
    elif tipo == "piscina":
        reps += _CONTRATANTE_IVONIR + _EXEC_SOBERANA

    return reps


# ---------------------------------------------------------------------------
# Substituição robusta em parágrafos (lida com runs fragmentados)
# ---------------------------------------------------------------------------

def _replace_in_para(para, old: str, new: str) -> bool:
    """Substitui texto em parágrafo, mesmo que fragmentado em vários runs."""
    runs = para.runs
    if not runs:
        return False

    full_text = ''.join(r.text for r in runs)
    if old not in full_text:
        return False

    # Mapa de posições de cada run no texto completo
    positions = []
    pos = 0
    for run in runs:
        ln = len(run.text)
        positions.append((pos, pos + ln))
        pos += ln

    match_start = full_text.index(old)
    match_end = match_start + len(old)

    # Primeiro e último runs envolvidos
    first_idx = next(i for i, (s, e) in enumerate(positions) if e > match_start)
    last_idx = next(i for i, (s, e) in enumerate(positions) if e >= match_end)

    # Constrói novo texto: mantém o que estava antes e depois do match
    before = full_text[positions[first_idx][0]:match_start]
    after = full_text[match_end:positions[last_idx][1]]
    runs[first_idx].text = before + new + after

    # Limpa os runs do meio e do último (que foram absorvidos)
    for i in range(first_idx + 1, last_idx + 1):
        runs[i].text = ''

    return True


def _process_doc_paras(doc, replacements: list[tuple]):
    """Aplica substituições em todos os parágrafos do documento (corpo + tabelas)."""
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        for old, new in replacements:
            if old in ''.join(r.text for r in para.runs):
                _replace_in_para(para, old, new)


# ---------------------------------------------------------------------------
# Fix 1: Remover highlight vermelho
# ---------------------------------------------------------------------------

def _remove_red_highlights(doc):
    """Remove highlight vermelho e sombreamento vermelho de todos os runs."""
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                continue
            # Remove w:highlight (destaque de texto)
            for hl in rPr.findall(qn('w:highlight')):
                rPr.remove(hl)
            # Remove w:shd com fill/color vermelho
            for shd in rPr.findall(qn('w:shd')):
                fill = shd.get(qn('w:fill'), '')
                color = shd.get(qn('w:color'), '')
                val = shd.get(qn('w:val'), '')
                # Remove qualquer shading vermelho (FF0000) ou highlight genérico
                if 'FF' in fill.upper() or 'FF' in color.upper() or fill.upper() in ('FF0000', 'RED'):
                    rPr.remove(shd)
                elif val in ('clear', 'nil') and fill.upper() in ('FF0000',):
                    rPr.remove(shd)


# ---------------------------------------------------------------------------
# Fix 2: Fundir parágrafo do logradouro com o do complemento (empreiteiro)
# ---------------------------------------------------------------------------

def _merge_address_paragraphs(doc):
    """
    Para o contrato empreiteiro: o logradouro e o complemento ficam em
    parágrafos separados. Funde os dois para evitar quebra de linha no
    endereço do contratante.
    """
    paras = doc.paragraphs
    for i, para in enumerate(paras[:-1]):
        if '{{ logradouro_contratante }}' in para.text:
            next_para = paras[i + 1]
            if '{{ complemento_contratante }}' in next_para.text:
                # Mover todos os w:r do próximo parágrafo para o atual
                para_elem = para._element
                next_elem = next_para._element
                run_tags = {qn('w:r'), qn('w:ins'), qn('w:del'), qn('w:hyperlink')}
                for child in list(next_elem):
                    if child.tag in run_tags:
                        para_elem.append(deepcopy(child))
                # Remover parágrafo do complemento
                next_elem.getparent().remove(next_elem)
                break


# ---------------------------------------------------------------------------
# Fix 3: Tabelas de serviços com tags Jinja2 para docxtpl
# ---------------------------------------------------------------------------

def _set_cell_text(tc_elem, text: str):
    """
    Define o texto de uma célula (w:tc), removendo todos os parágrafos extras
    e preservando o estilo do primeiro run encontrado.
    """
    all_paras = tc_elem.findall(qn('w:p'))
    if not all_paras:
        return

    # Coletar rPr do primeiro run de qualquer parágrafo (para preservar estilo)
    rPr_clone = None
    for ap in all_paras:
        existing_run = ap.find(qn('w:r'))
        if existing_run is not None:
            rPr_existing = existing_run.find(qn('w:rPr'))
            if rPr_existing is not None:
                rPr_clone = deepcopy(rPr_existing)
                # Limpar highlights
                for hl in rPr_clone.findall(qn('w:highlight')):
                    rPr_clone.remove(hl)
                for shd in rPr_clone.findall(qn('w:shd')):
                    rPr_clone.remove(shd)
            break

    # Usar o primeiro parágrafo e remover os demais
    p = all_paras[0]
    for extra_p in all_paras[1:]:
        tc_elem.remove(extra_p)

    # Limpar todos os runs e elementos inline do parágrafo mantido
    run_like_tags = {qn('w:r'), qn('w:hyperlink'), qn('w:bookmarkStart'),
                     qn('w:bookmarkEnd'), qn('w:ins'), qn('w:del')}
    for child in list(p):
        if child.tag in run_like_tags:
            p.remove(child)

    if not text:
        return

    # Criar novo run com o texto
    r_elem = etree.SubElement(p, qn('w:r'))
    if rPr_clone is not None:
        r_elem.insert(0, rPr_clone)

    t_elem = etree.SubElement(r_elem, qn('w:t'))
    t_elem.text = text
    if text != text.strip():
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _setup_tabelas_servicos(doc):
    """
    Para todas as tabelas do documento que contenham dados de serviços:
    - Mantém apenas a linha de cabeçalho
    - Insere 3 linhas docxtpl: {%tr for %} | linha de dados | {%tr endfor %}

    Em docxtpl as linhas de controle ({%tr for/endfor %}) são removidas do
    output e substituídas pelos blocos Jinja2 correspondentes.  A linha
    de dados (entre elas) é que efetivamente se repete.
    """
    for tbl_idx, tbl in enumerate(doc.tables):
        rows = tbl.rows
        if len(rows) < 2:
            continue

        num_cols = len(rows[0].cells)

        # Só processar tabelas de serviços (primeira cel do cabeçalho = "Item")
        header_text = rows[0].cells[0].text.strip().lower()
        if 'item' not in header_text:
            continue

        # Clonar o estilo da primeira linha de dados
        template_tr = deepcopy(rows[1]._tr)

        # Remover todas as linhas de dados (manter só cabeçalho)
        tbl_elem = tbl._tbl
        for row in rows[1:]:
            tbl_elem.remove(row._tr)

        # ── 1. Linha de controle FOR (toda a linha será removida pelo docxtpl) ──
        for_ctrl = deepcopy(template_tr)
        for_cells = for_ctrl.findall(qn('w:tc'))
        _set_cell_text(for_cells[0], '{%tr for s in servicos %}')
        for c in for_cells[1:]:
            _set_cell_text(c, '')
        tbl_elem.append(for_ctrl)

        # ── 2. Linha de dados (repete para cada serviço) ─────────────────────
        data_row = deepcopy(template_tr)
        data_cells = data_row.findall(qn('w:tc'))

        if num_cols == 2:
            _set_cell_text(data_cells[0], '{{ s.num }}')
            if len(data_cells) > 1:
                _set_cell_text(data_cells[1], '{{ s.descricao }}')
        elif num_cols >= 5:
            _set_cell_text(data_cells[0], '{{ s.num }}')
            if len(data_cells) > 1:
                _set_cell_text(data_cells[1], '{{ s.descricao }}')
            if len(data_cells) > 2:
                _set_cell_text(data_cells[2], '{{ s.unidade }}')
            if len(data_cells) > 3:
                _set_cell_text(data_cells[3], '{{ s.valor }}')
            if len(data_cells) > 4:
                _set_cell_text(data_cells[4], '{{ s.prazo }}')
        else:
            _set_cell_text(data_cells[0], '{{ s.num }}')
            for c in data_cells[1:]:
                _set_cell_text(c, '')

        tbl_elem.append(data_row)

        # ── 3. Linha de controle ENDFOR ───────────────────────────────────────
        end_ctrl = deepcopy(template_tr)
        end_cells = end_ctrl.findall(qn('w:tc'))
        _set_cell_text(end_cells[0], '{%tr endfor %}')
        for c in end_cells[1:]:
            _set_cell_text(c, '')
        tbl_elem.append(end_ctrl)


# ---------------------------------------------------------------------------
# Preparação de templates
# ---------------------------------------------------------------------------

def preparar_template(tipo: str, force: bool = False) -> Path:
    info = CONTRATOS[tipo]
    template_path = TEMPLATES_DIR / info["template"]

    if template_path.exists() and not force:
        return template_path

    modelo_path = MODELOS_DIR / info["modelo"]
    if not modelo_path.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    doc = Document(str(modelo_path))

    # 1. Substituições de texto (variáveis Jinja2)
    replacements = _get_replacements(tipo)
    _process_doc_paras(doc, replacements)

    # 2. Para empreiteiro: fundir parágrafo do logradouro com o do complemento
    if tipo == "empreiteiro":
        _merge_address_paragraphs(doc)

    # 3. Remover highlights/sombreamento vermelho
    _remove_red_highlights(doc)

    # 4. Configurar tabelas de serviços com tags docxtpl
    _setup_tabelas_servicos(doc)

    doc.save(str(template_path))
    return template_path


def preparar_todos_templates(force: bool = False):
    for tipo in CONTRATOS:
        try:
            preparar_template(tipo, force=force)
            print(f"[OK] Template '{tipo}' pronto.")
        except FileNotFoundError as e:
            print(f"[AVISO] {e}")
        except Exception as e:
            import traceback
            print(f"[ERRO] Template '{tipo}': {e}")
            traceback.print_exc()


# ---------------------------------------------------------------------------
# Geração do contrato preenchido
# ---------------------------------------------------------------------------

def gerar_contrato(tipo: str, dados: dict, servicos: list[dict]) -> Path:
    template_path = preparar_template(tipo)
    tpl = DocxTemplate(str(template_path))

    for i, s in enumerate(servicos, 1):
        s["num"] = i

    context = {**dados, "servicos": servicos}
    tpl.render(context)

    nome_safe = re.sub(r"[^\w]", "_", dados.get("nome_contratante", "contrato"))
    filename = f"{tipo}_{nome_safe}.docx"
    out_path = GENERATED_DIR / filename
    tpl.save(str(out_path))
    return out_path


def listar_tipos_contrato() -> list[dict]:
    import database as _db
    resultado = [{"id": k, "label": v["label"], "origem": "sistema"} for k, v in CONTRATOS.items()]
    try:
        extras = _db.listar_tipos_contrato_db()
        for t in extras:
            if t["key"] not in CONTRATOS:
                resultado.append({"id": t["key"], "label": t["label"], "origem": "personalizado"})
    except Exception:
        pass
    return resultado


def processar_template_vermelho(modelo_bytes: bytes) -> bytes:
    """
    Converte um .docx onde textos em VERMELHO são os dados que variam por contrato.
    O usuário pinta o texto real (ex: "Cassio Antônio da Silva" em vermelho).
    A IA identifica pelo contexto qual variável Jinja2 cada texto representa.
    Tabelas com header "Item/Descrição/Unidade/Valor/Prazo" são convertidas
    automaticamente para loop docxtpl ({%tr for s in servicos %}).
    """
    import tempfile, os, json as _json
    from groq import Groq as _Groq

    _REDS = {"FF0000", "C00000", "A50021", "FF0001"}

    def _is_red(run) -> bool:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            return False
        c = rPr.find(qn("w:color"))
        if c is not None:
            val = c.get(qn("w:val"), "").upper()
            if val in _REDS or (val.startswith("FF") and val != "FFFFFF"):
                return True
        hl = rPr.find(qn("w:highlight"))
        if hl is not None and hl.get(qn("w:val"), "").lower() == "red":
            return True
        return False

    def _remove_red(run):
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            return
        for el in list(rPr.findall(qn("w:color"))): rPr.remove(el)
        for el in list(rPr.findall(qn("w:highlight"))): rPr.remove(el)

    def _para_red_text(para) -> str:
        return "".join(r.text for r in para.runs if _is_red(r)).strip()

    def _para_full_text(para) -> str:
        return "".join(r.text for r in para.runs).strip()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(modelo_bytes)
        tmp_path = tmp.name

    doc = Document(tmp_path)

    # ── 1. Identificar tabelas de serviços e coletar amostras de texto vermelho ──
    service_table_idxs = set()
    samples = []  # [{red_text, context, is_upper}]

    body_paras = list(doc.paragraphs)

    for i, para in enumerate(body_paras):
        red = _para_red_text(para)
        if not red:
            continue
        ctx_prev = _para_full_text(body_paras[i - 1]) if i > 0 else ""
        ctx_next = _para_full_text(body_paras[i + 1]) if i < len(body_paras) - 1 else ""
        full = _para_full_text(para)
        samples.append({
            "red_text": red,
            "context": f"{ctx_prev} | {full} | {ctx_next}",
            "is_upper": red == red.upper() and red != red.lower(),
        })

    for tbl_idx, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue
        header = " ".join(c.text.lower() for c in tbl.rows[0].cells)
        if any(w in header for w in ("item", "descriç", "serviç", "unidade")):
            service_table_idxs.add(tbl_idx)
            continue
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    red = _para_red_text(para)
                    if not red:
                        continue
                    full = _para_full_text(para)
                    samples.append({
                        "red_text": red,
                        "context": full,
                        "is_upper": red == red.upper() and red != red.lower(),
                    })

    # ── 2. Pedir ao Groq o mapeamento texto → variável ──
    VARIAVEIS = (
        "nome_contratante: nome completo do contratante (cliente, pessoa física)\n"
        "cpf_contratante: CPF do contratante (formato 000.000.000-00)\n"
        "rg_contratante: número do RG do contratante\n"
        "uf_rg_contratante: estado emissor do RG (ex: GO)\n"
        "nacionalidade_contratante: nacionalidade (ex: brasileiro(a))\n"
        "estado_civil_contratante: estado civil (ex: solteiro(a), casado(a))\n"
        "profissao_contratante: profissão do contratante\n"
        "logradouro_contratante: rua e número do endereço do contratante\n"
        "complemento_contratante: complemento do endereço do contratante\n"
        "nome_executora: nome da empresa ou pessoa executora\n"
        "cnpj_executora: CNPJ da executora (formato 00.000.000/0000-00)\n"
        "cpf_executora: CPF da executora se pessoa física\n"
        "endereco_executora: endereço da executora\n"
        "data_inicio: data de início da obra (DD/MM/AAAA)\n"
        "prazo_execucao: prazo de execução (ex: 6 semanas, 3 meses)\n"
        "data_conclusao: data prevista de conclusão (DD/MM/AAAA)\n"
        "valor_multa_fmt: valor da multa em R$ (ex: R$ 10.000,00)\n"
        "valor_multa_extenso: valor da multa por extenso (ex: dez mil reais)\n"
        "cidade_assinatura: cidade onde o contrato é assinado (ex: Anápolis-GO)\n"
        "data_assinatura: data de assinatura por extenso (ex: 22 de Janeiro de 2026)\n"
        "nome_contratante | upper: nome do contratante em MAIÚSCULAS\n"
        "nome_executora | upper: nome da executora em MAIÚSCULAS\n"
        "cnpj_executora (usado em maiúsculas junto com nome): use nome_executora | upper – cnpj_executora\n"
    )

    mapping = {}
    if samples:
        prompt = (
            "Você analisa contratos de construção civil brasileiros.\n"
            "Os textos abaixo foram marcados em VERMELHO no contrato — são dados que variam por contrato.\n"
            "Mapeie cada texto vermelho para a variável correta da lista.\n\n"
            f"VARIÁVEIS DISPONÍVEIS:\n{VARIAVEIS}\n\n"
            "TEXTOS VERMELHOS COM CONTEXTO:\n"
            + _json.dumps(samples, ensure_ascii=False, indent=2)
            + "\n\nRetorne APENAS JSON: {\"texto_vermelho\": \"nome_variavel\", ...}\n"
            "Use exatamente o texto vermelho como chave. "
            "Se aparecer em MAIÚSCULAS, use a variável com '| upper'. "
            "Se não souber, use 'DESCONHECIDO'."
        )

        try:
            client_groq = _Groq(api_key=os.environ.get("GROQ_API_KEY"))
            resp = client_groq.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=2048,
            )
            raw = resp.choices[0].message.content.strip()
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw)
            mapping = _json.loads(raw)
        except Exception:
            mapping = {}

    # ── 3. Aplicar mapeamento no documento ──
    def _apply_para(para):
        red_runs = [r for r in para.runs if _is_red(r)]
        if not red_runs:
            return
        red_text = "".join(r.text for r in red_runs).strip()
        var_name = mapping.get(red_text, red_text)
        replacement = "{{ " + var_name + " }}"
        red_runs[0].text = replacement
        _remove_red(red_runs[0])
        for r in red_runs[1:]:
            r.text = ""
            _remove_red(r)

    for para in doc.paragraphs:
        _apply_para(para)

    for tbl_idx, tbl in enumerate(doc.tables):
        if tbl_idx in service_table_idxs:
            continue  # tratada pelo _setup_tabelas_servicos abaixo
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_para(para)

    # ── 4. Processar tabelas de serviços (loop docxtpl) ──
    _setup_tabelas_servicos(doc)
    _remove_red_highlights(doc)

    out_path = tmp_path.replace(".docx", "_processed.docx")
    doc.save(out_path)
    with open(out_path, "rb") as f:
        result = f.read()

    try:
        os.unlink(tmp_path)
        os.unlink(out_path)
    except Exception:
        pass

    return result


def gerar_contrato_dinamico(key: str, modelo_bytes: bytes, dados: dict, servicos: list[dict]) -> Path:
    """Gera contrato a partir de template .docx armazenado no banco (já com variáveis Jinja2)."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(modelo_bytes)
        tmp_path = tmp.name

    tpl = DocxTemplate(tmp_path)
    for i, s in enumerate(servicos, 1):
        s["num"] = i
    context = {**dados, "servicos": servicos}
    tpl.render(context)

    nome_safe = re.sub(r"[^\w]", "_", dados.get("nome_contratante", "contrato"))
    filename = f"{key}_{nome_safe}.docx"
    out_path = GENERATED_DIR / filename
    tpl.save(str(out_path))
    return out_path
